#!/usr/bin/env python3
"""
Export Utilities — DOCX generation from Markdown
===================================================
Converts LLM-generated markdown resume/cover letter to a properly
formatted Word document. No cloud APIs, runs fully locally.

Requires: python-docx  (pip install python-docx)
"""

import io
import re
from typing import Optional


def markdown_to_docx_bytes(markdown_text: str, title: str = "Document") -> bytes:
    """
    Convert a markdown string to a .docx file returned as bytes.

    Handles:
      - # H1  ## H2  ### H3  headings
      - **bold** inline
      - - / * bullet lists
      - --- horizontal rules (rendered as a paragraph of dashes)
      - Plain paragraphs

    Returns raw bytes suitable for st.download_button.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        )

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- Default body font ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def _add_horizontal_rule(document):
        p = document.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _apply_inline_bold(paragraph, text: str):
        """Add a run with **bold** segments parsed inline."""
        parts = re.split(r"\*\*(.+?)\*\*", text)
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            run.bold = (i % 2 == 1)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_bold(p, line[2:])
        elif re.match(r"^[-_]{3,}$", line.strip()):
            _add_horizontal_rule(doc)
        elif line.strip():
            p = doc.add_paragraph()
            _apply_inline_bold(p, line)
        else:
            doc.add_paragraph("")

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
